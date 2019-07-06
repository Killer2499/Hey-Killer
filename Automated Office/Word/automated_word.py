from docx import Document
from docx.shared import Inches

document=Document()


title=raw_input("Title of the Document: ")
document.add_heading(title,0)

def add_paragraph():
    paragraph_data=raw_input("Paragraph Data:")
    paragraph=document.add_paragraph(paragraph_data)


def add_heading():
    heading_data=raw_input("Header Data:")
    heading=document.add_heading(heading_data,level=1)

def add_picture():    
    picture_path=raw_input("Image Path:")
    specifications=raw_input("Image Dimension Specifications or Default")
    if specifications=="Yes":
        height=raw_input("Height in Inches:")
        width=raw_input("Width in Inches:")
    else:
        height=4;
        width=4;
    picture=document.add_picture(picture_path,width=Inches(width),height=Inches(height))

#def table():
#    table_dimensions=raw_input("Rows,Columns:")
#    table_rows=table_dimensions.split(',')[0]
#    table_cols=table_dimensions.split(',')[1]
#    table=document.add_table(rows=table_rows,cols=table_cols)
#   hdr_cells=table
flag=True
while flag==True:
    data_type=input("Enter Type to add: 1.Paragraph 2.Heading 3.Picture")
    if (data_type==1):
        add_paragraph()
    elif(data_type==2):
        add_heading()
    elif (data_type==3):
        add_picture()
    else:
        flag=False
document.add_page_break()
document_location=("FirstDoc.docx")
document.save(document_location)
