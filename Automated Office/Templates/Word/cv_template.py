from docx import Document
from docx.shared import Inches
from datetime import date
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH,WD_BREAK
from docx.enum.style import WD_STYLE
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Pt
import sys
sys.path.append('../../../')
from speak import say
  
def cv():
    document = Document()

    paragraph_styles = document.styles
    paragraph_styles = paragraph_styles.add_style('ParagraphStyle',WD_STYLE_TYPE.PARAGRAPH)
    paragraph_font = paragraph_styles.font
    paragraph_font.size = Pt(12)
    paragraph_font.name = 'Segoe UI Historic'

    title_styles=document.styles
    title_styles = title_styles.add_style('TitleStyle',WD_STYLE_TYPE.CHARACTER)
    title_font=title_styles.font
    title_font.color.rgb=RGBColor(247,93,93)
    title_font.bold=True
    title_font.size=Pt(25)
    title_font.name='Segoe UI Historic'

    heading_styles=document.styles
    heading_styles = heading_styles.add_style('HeadingStyle',WD_STYLE_TYPE.CHARACTER)
    heading_font=heading_styles.font
    heading_font.color.rgb=RGBColor(247,93,93)
    heading_font.bold=True
    heading_font.size=Pt(17)
    heading_font.name='Segoe UI Historic'

    say("What's your name")
    name=input("Name:")
    #name='Sanath Singavarapu'
    name_title=document.add_paragraph()
    name_title=name_title.add_run(name,style="TitleStyle")
    name_title.size=Pt(40)

    
    say("Hey"+name.split(' ')[0]+"What's your Email")   
    email=input("Email:")
    say("Can I have ur Number:")
    phone=input("Phone:")
    #email='sanathsingavarapu99@gmail.com'
    #phone='+91-8919142764'
   
    
    contact_info="Email: "+email+" | Telephone: "+phone
    contact=document.add_paragraph(contact_info,style='ParagraphStyle')

    spacing = contact.add_run()
    spacing.add_break(WD_BREAK.LINE)

    say("State your Personal Statement")
    statement_info=input("Personal Statement:")
    #statement_info='A committed,knowledgeable and capable Reasearch Fellow.Highly experienced in project and team management, strategic planning and budget management. A confident presenter able to impart complex information to audiences of all levels.'
    statement_header=document.add_paragraph().add_run("Personal Statement",style="HeadingStyle")    
    statement=document.add_paragraph(statement_info,style='ParagraphStyle')

    say("List your Technical Skills")
    #skills_info=input(" Technical Skills:")
    skills_info='Html,Css,Js,Jquery,Php,Mysql,T-sql,U-sql,Python,C,C++,Bootstrap,PL/SQL,NO-SQL,Power Bi,Azure Data Factory,Azure Data Bricks,React Native,React Js,Arduino,Raspberry Pi,Adobe Photoshop'
    skills_header=document.add_paragraph().add_run("Technical Skills",style="HeadingStyle")
    skills=document.add_paragraph(skills_info,style='ParagraphStyle')

    def add_experience(experience_count):
        if (experience_count==0):
            experience_header=document.add_paragraph().add_run("Experience",style="HeadingStyle")

        say("Name of the Company")
        company_info=input("Company Name:")
        experience=document.add_paragraph()
        company=experience.add_run(company_info)
        company_font=company.font
        company_font.bold=True
        company_font.size=Pt(15)
        company_font.name='Segoe UI Historic'
        
        spacing = experience.add_run()
        spacing.add_break(WD_BREAK.LINE)

        say("Your Role in the company")
        role_info=input("Role:")
        role=experience.add_run(role_info)
        
        spacing = experience.add_run()
        spacing.add_break(WD_BREAK.LINE)

        say("How long")
        time_period_info=input("Time Period:")
        time_period=experience.add_run(time_period_info)

        experience_count+=1
        if(experience_count>=1):
            say("Wanna Add one more")
            option=input("Wanna Add More(y/n):")
            if option=='y' or option=='yes':
                add_experience(1)
            else:
                pass
    add_experience(0)

    print ("---Education---")

    def add_education(education_count):
        if (education_count==0):
            education_header=document.add_paragraph().add_run("Education",style="HeadingStyle")

        
        say("College's name")
        college_info=input("College Name:")
        education=document.add_paragraph()
        college=education.add_run(college_info)
        college_font=college.font
        college_font.bold=True
        college_font.size=Pt(15)
        college_font.name='Segoe UI Historic'
        
        
        spacing = education.add_run()
        spacing.add_break(WD_BREAK.LINE)

        say("Time Period")
        time_period_info=input("Time Period:")
        time_period=education.add_run(time_period_info)

        education_count+=1
        if(education_count>=1):
            
            say("Want to Add more")  
            option=input("Wanna Add More(y/n):")
            if option=='y' or option=='yes':
                add_education(1)
            else:
                pass
    add_education(0)

    document.add_page_break()

    print ("---Project---")

    def add_project(project_count):
        if (project_count==0):
            project_header=document.add_paragraph().add_run("Projects",style="HeadingStyle")

        say("Enter Project Name and Details")
        project_title_info=input("Project Name:")
        project_details_info=input("Project Details:")

        
        projects=document.add_paragraph()
        project_title=projects.add_run(project_title_info)
        project_title_font=project_title.font
        project_title_font.size=Pt(12)
        project_title_font.bold=True
        project_title_font.name='Segoe UI Historic'

        project_details=projects.add_run(" : " +project_details_info)
        project_details_font=project_details.font
        project_details_font.size=Pt(12)
        project_details_font.name='Segoe UI Historic'

        project_count+=1
        if(project_count>=1):
            option=input("Wanna Add More(y/n):")
            if option=='y' or option=='yes':
                add_project(1)
            else:
                pass
    add_project(0)
    
    
    print ("---Awards---")

    def add_awards(award_count):
        if (award_count==0):
            award_header=document.add_paragraph().add_run("Awards",style="HeadingStyle")

        say("Mention about your awards")
        competition_info=input("Competition Name:")
        winning_info=input("Position:")

        
        awards=document.add_paragraph()
        competition=awards.add_run(competition_info)
        competition_font=competition.font
        competition_font.size=Pt(12)
        competition_font.bold=True
        competition_font.name='Segoe UI Historic'

        winning=awards.add_run(" : " +winning_info)
        winning_font=winning.font
        winning_font.size=Pt(12)
        winning_font.name='Segoe UI Historic'

        award_count+=1
        if(award_count>=1):
            option=input("Wanna Add More(y/n):")
            if option=='y' or option=='yes':
                add_awards(1)
            else:
                pass
    add_awards(0)

    say("Where Should I save this")
    location=input("Location:")
    document.save(location)

cv()

cv()
