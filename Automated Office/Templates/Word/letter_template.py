from docx import Document
from docx.shared import Inches
from datetime import date
from docx.enum.text import WD_ALIGN_PARAGRAPH,WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Pt
import sys
sys.path.append('../../../')
from speak import say

def letter():
    document=Document()

    paragraph_styles = document.styles
    paragraph_styles = paragraph_styles.add_style('ParagraphStyle',WD_STYLE_TYPE.PARAGRAPH)
    paragraph_font = paragraph_styles.font
    paragraph_font.size = Pt(12)
    paragraph_font.name = 'Times New Roman'

    say("Enter from address:")
    #From Address
    from_input=input("From address:")
    from_split=from_input.split(',')

    data=''
    for i in range(0,len(from_split)):
        if(i==len(from_split)-1):
            data=data+from_split[i]+"."
        else:
            data=data+from_split[i]+",\n"


    #from_input=input("From address:")

    #from_split=from_input.split(',')
    #from_name=from_split[0]

    #data=''
    #for i in range(1,len(from_split)):
    #   data=data+from_split[i]+" . "
        
    #data=data[:-2]


    #from_title_name=document.add_paragraph()
    #from_title_name.add_run(from_name, style = 'CommentsStyle').bold = True
    #from_title_name.bold=True
    #from_title_name.alignment =1 #0-L ,1-C, 2-R
    
    
    from_address=document.add_paragraph(data,style='ParagraphStyle')
    from_address.name = 'Times New Roman'
    from_address.size = Pt(14)
    #from_address.alignment =1
   
    spacing = from_address.add_run()
    spacing.add_break(WD_BREAK.LINE)

    #Date
    date_today=str(date.today())
    year=int(date_today.split('-')[0])
    month=int(date_today.split('-')[1])
    day=int(date_today.split('-')[2])
    date_words=date(day=day, month=month, year=year).strftime('%d %B %Y')
    letter_date=document.add_paragraph(date_words,style='ParagraphStyle')

    spacing = letter_date.add_run()
    spacing.add_break(WD_BREAK.LINE)

    say("To whom should I send the letter")
    #To Address
    to_input=input("To address:")
    to_split=to_input.split(',')

    data=''
    for i in range(0,len(to_split)):
        if(i==len(to_split)-1):
            data=data+to_split[i]+"."
        else:
            data=data+to_split[i]+",\n"

    to_address=document.add_paragraph(data,style='ParagraphStyle')

    spacing = to_address.add_run()
    spacing.add_break(WD_BREAK.LINE)

    #Greetings
    say("Give your Greetings")
    greeting_input=input("Greetings:")
    greetings=document.add_paragraph(greeting_input,style='ParagraphStyle')

    #Body
    say("Enter the content of the letter")
    body_input=input("Body of the Letter:")
    body_input=body_input.split('<new_paragraph>')
    for i in range(0,len(body_input)):
        body=document.add_paragraph(body_input[i],style='ParagraphStyle')

    spacing = body.add_run()
    spacing.add_break(WD_BREAK.LINE)

    #Closing
    say("Closing Greetings")
    closing_input=input("Closing Greeting:")
    closing=document.add_paragraph(closing_input,style='ParagraphStyle')


    #Signature
    say("Wanna Add your Signature")
    signature_image_option=input("Wanna Add Signature Image(Yes/No):")
    if(signature_image_option=='Yes' or signature_image_option=='yes'):
        signature_image_path=input("Image Path:")
        document.add_picture(signature_image_path,width=Inches(2),height=Inches(1))
    else:
        pass

    #Name and Designation
    say("Name and Designation")
    designation_details=input("Name and Designation:")
    designation_details=designation_details.split(',')
    data=''
    for i in range(0,len(designation_details)):
        if(i==len(designation_details)-1):
            data=data+designation_details[i]+" "
        else:
            data=data+designation_details[i]+",\n"
      
    designation=document.add_paragraph(data,style='ParagraphStyle')
    document.add_page_break()

    #Saving Document
    say("Where should I save this")
    location=input("Save Document as")
    #location="C:\Users\Sanath\Desktop\Hey Killer!\Automated Office\Templates\Word\Sample_Letter.docx"
    document.save(location)


letter()

